from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO
import io


@dataclass
class LoadedDocument:
    name: str
    text: str
    source_type: str


def _read_pdf(file_obj: BinaryIO) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf. Install requirements.txt first.") from exc

    file_bytes = file_obj.read()
    file_obj = io.BytesIO(file_bytes)
    reader = PdfReader(file_obj)
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if len(text.split()) < 12:
            ocr_text = _ocr_pdf_page(file_bytes, page_number - 1)
            if ocr_text:
                text = f"{text}\n{ocr_text}".strip()
        if text.strip():
            pages.append(f"[Page {page_number}]\n{text}")
    return "\n\n".join(pages)


def _ocr_pdf_page(file_bytes: bytes, zero_based_page_number: int) -> str:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return ""

    try:
        pdf = pdfium.PdfDocument(file_bytes)
        page = pdf[zero_based_page_number]
        bitmap = page.render(scale=2)
        image = bitmap.to_pil()
    except Exception:
        return ""
    return _ocr_image_object(image)


def _ocr_image_object(image: object) -> str:
    try:
        import pytesseract
    except ImportError:
        return ""

    try:
        text = pytesseract.image_to_string(image)
    except Exception:
        return ""
    return text.strip()


def _read_image(file_obj: BinaryIO) -> str:
    try:
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("Image upload requires Pillow. Install requirements.txt first.") from exc

    image = Image.open(file_obj)
    text = _ocr_image_object(image)
    if not text:
        raise RuntimeError(
            "No readable text could be extracted from this image. "
            "OCR support may be unavailable, or the screenshot may be too unclear."
        )
    return f"[Image OCR]\n{text}"


def _read_docx(file_obj: BinaryIO) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Word document support requires python-docx. Install requirements.txt first.") from exc

    document = Document(file_obj)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _read_text(file_obj: BinaryIO) -> str:
    raw = file_obj.read()
    if isinstance(raw, str):
        return raw
    return raw.decode("utf-8", errors="ignore")


def load_uploaded_file(file_obj: BinaryIO, filename: str) -> LoadedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(file_obj)
        source_type = "pdf"
    elif suffix == ".docx":
        text = _read_docx(file_obj)
        source_type = "docx"
    elif suffix in {".png", ".jpg", ".jpeg"}:
        text = _read_image(file_obj)
        source_type = "image"
    elif suffix in {".txt", ".md"}:
        text = _read_text(file_obj)
        source_type = "text"
    else:
        raise ValueError("Unsupported file type. Please upload a PDF, DOCX, image, TXT, or MD file.")

    return LoadedDocument(name=filename, text=normalise_text(text), source_type=source_type)


def load_sample_document(path: str | Path) -> LoadedDocument:
    path = Path(path)
    text = path.read_text(encoding="utf-8")
    return LoadedDocument(name=path.name, text=normalise_text(text), source_type=path.suffix.lstrip("."))


def normalise_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").split("\n")]
    cleaned = []
    blank_seen = False
    for line in lines:
        if not line:
            if not blank_seen:
                cleaned.append("")
            blank_seen = True
            continue
        cleaned.append(line)
        blank_seen = False
    return "\n".join(cleaned).strip()
